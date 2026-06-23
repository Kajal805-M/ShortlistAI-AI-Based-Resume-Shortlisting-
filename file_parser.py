from io import BytesIO

from docx import Document
from pypdf import PdfReader


class DocumentParseError(Exception):
    pass


def extract_text(file_stream, extension):
    payload = file_stream.read()
    if not payload:
        raise DocumentParseError("The uploaded file is empty.")

    try:
        if extension == ".pdf":
            return _extract_pdf(payload)
        if extension == ".docx":
            return _extract_docx(payload)
    except DocumentParseError:
        raise
    except Exception as error:
        raise DocumentParseError("The document appears damaged or password-protected.") from error

    raise DocumentParseError("Unsupported document format.")


def _extract_pdf(payload):
    reader = PdfReader(BytesIO(payload))
    if reader.is_encrypted:
        raise DocumentParseError("Password-protected PDFs are not supported.")
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(payload):
    document = Document(BytesIO(payload))
    content = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            content.extend(cell.text for cell in row.cells)
    return "\n".join(content)

